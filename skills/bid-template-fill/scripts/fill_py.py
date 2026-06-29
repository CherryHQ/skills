#!/usr/bin/env python3
"""DOCX template fill engine v2.2"""
import json, sys, os, re
from datetime import date

if len(sys.argv) < 3:
    print('Usage: python fill_py.py <template.docx> <output.docx> [data.json]')
    sys.exit(1)

template_path = sys.argv[1]
output_path = sys.argv[2]
data_path = sys.argv[3] if len(sys.argv) > 3 else None

if not os.path.exists(template_path):
    print(f'[ERROR] Template not found: {template_path}')
    sys.exit(1)

replacements = {}
if data_path:
    try:
        with open(data_path, 'r', encoding='utf-8-sig') as f:
            replacements = json.load(f)
    except json.JSONDecodeError as e:
        print(f'[ERROR] JSON: {e}')
        sys.exit(1)
    print(f'  Loaded {len(replacements)} keys')

today = date.today()
today_str = f'{today.year} 年 {today.month:02d} 月 {today.day:02d} 日'

try:
    from docx import Document
except ImportError:
    print("[ERROR] Need python-docx. Run: pip install python-docx")
    sys.exit(1)

print(f'1. Loading: {template_path}')
doc = Document(template_path)
total = 0
all_placeholders_found = set()

DATE_PATTERN = re.compile(r'(20\d{2})[_\s]*年[_\s]*(\d{1,2})[_\s]*月[_\s]*(\d{1,2})[_\s]*日')
DATE_EMPTY_PATTERN = re.compile(r'_*(20\d{2})[_\s]*年[_\s]*月[_\s]*日')
DATE_BARE_PATTERN = re.compile(r'_*年[_\s]*月[_\s]*日')

def merge_runs_text(runs):
    return "".join(r.text for r in runs)

def rebuild_para_text(para, new_text):
    global total
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    total += 1

def replace_dates_in_para(para):
    """Phase 1: Replace all date patterns with today"""
    runs = para.runs
    full = merge_runs_text(runs)

    def replacer(m):
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if y == today.year and mo == today.month and d == today.day:
            return m.group(0)
        return today_str

    changed = False
    new_text, c = DATE_PATTERN.subn(replacer, full)
    if c > 0:
        full = new_text
        changed = True
    new_text, c2 = DATE_EMPTY_PATTERN.subn(today_str, full)
    if c2 > 0:
        full = new_text
        changed = True
    new_text, c3 = DATE_BARE_PATTERN.subn(today_str, full)
    if c3 > 0:
        full = new_text
        changed = True
    if changed:
        rebuild_para_text(para, full)

def replace_placeholders_in_para(para):
    """Phase 2: Replace all placeholders with mapped values"""
    runs = para.runs
    full = merge_runs_text(runs)
    changed = False
    for key, value in sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True):
        if key not in full:
            continue
        full = full.replace(key, value)
        changed = True
    if changed:
        rebuild_para_text(para, full)

def scan_placeholders(para):
    """Scan for remaining placeholders after fill"""
    full = merge_runs_text(para.runs)
    for pat in [r"_[^_]+_", r"\[[^\]]+\]", r"（[^）]+）"]:
        for m in re.findall(pat, full):
            if not m.startswith("[Content") and not m.startswith("[ISO"):
                all_placeholders_found.add(m)


def process_all(doc):
    """Process all paragraphs, tables, headers, footers.
    Phase 1: dates -> Phase 2: placeholders -> Phase 3: scan"""
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for section in doc.sections:
        all_paras.extend(section.header.paragraphs)
        all_paras.extend(section.footer.paragraphs)

    # Phase 1: Dates first (before placeholder values can interfere)
    for para in all_paras:
        replace_dates_in_para(para)

    # Phase 2: Placeholders
    for para in all_paras:
        replace_placeholders_in_para(para)

    # Phase 3: Scan for remaining
    for para in all_paras:
        scan_placeholders(para)

print("2. Phase 1: Date replacement...")
print("3. Phase 2: Placeholder replacement...")
process_all(doc)
print(f"  Processed {total} occurrences")

if all_placeholders_found:
    filled_keys = set(replacements.keys())
    unfilled = all_placeholders_found - filled_keys
    if unfilled:
        print(f"\n  [WARN] {len(unfilled)} unmatched placeholders:")
        for p in sorted(unfilled)[:20]:
            print(f"    - {p}")
        if len(unfilled) > 20:
            print(f"    ... and {len(unfilled) - 20} more")
        print("  Check if replacement data includes these keys.")

print(f"4. Saving: {output_path}")
output_dir = os.path.dirname(os.path.abspath(output_path))
if output_dir and not os.path.exists(output_dir):
    os.makedirs(output_dir, exist_ok=True)

final_path = output_path
if os.path.exists(output_path):
    try:
        os.remove(output_path)
    except OSError:
        dirn = os.path.dirname(output_path)
        name, ext = os.path.splitext(os.path.basename(output_path))
        n = 2
        while os.path.exists(os.path.join(dirn, f"{name}-V{n}{ext}")):
            n += 1
        final_path = os.path.join(dirn, f"{name}-V{n}{ext}")
        print(f"  Locked => {os.path.basename(final_path)}")


try:
    doc.save(final_path)

    # Phase 4: Direct patch header/footer XML
    # python-docx only accesses default header; we patch header1/2/3.xml too
    import tempfile as _tf, zipfile as _zf, shutil as _sh
    _tmpdir = _tf.mkdtemp(prefix=".hdr_")
    try:
        with _zf.ZipFile(final_path, "r") as _z:
            _z.extractall(_tmpdir)
        _patched = 0

        def _dr_py(_m):
            _y, _mo, _d = int(_m.group(1)), int(_m.group(2)), int(_m.group(3))
            if _y == today.year and _mo == today.month and _d == today.day:
                return _m.group(0)
            return today_str

        for _root, _dirs, _files in os.walk(_tmpdir):
            for _fn in _files:
                if ("header" in _fn.lower() or "footer" in _fn.lower()) and _fn.endswith(".xml"):
                    _fp = os.path.join(_root, _fn)
                    with open(_fp, "r", encoding="utf-8") as _fh:
                        _hc = _fh.read()
                    _changed = False

                    # 1) Placeholder replace: direct string replace (preserves XML)
                    for _k, _v in sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True):
                        if _k in _hc:
                            _esc = str(_v).replace("&", "&amp;").replace("<", "&lt;")
                            _esc = _esc.replace(">", "&gt;").replace('"', "&quot;")
                            _hc = _hc.replace(_k, _esc)
                            _changed = True

                    # 2) Date replace: per <w:t> element, preserving XML structure
                    def _rdwt(_wm):
                        _ts, _tx, _te = _wm.group(1), _wm.group(2), _wm.group(3)
                        _ot = _tx
                        _tx, _ = DATE_PATTERN.subn(_dr_py, _tx)
                        _tx, _ = DATE_EMPTY_PATTERN.subn(today_str, _tx)
                        _tx, _ = DATE_BARE_PATTERN.subn(today_str, _tx)
                        if _tx != _ot:
                            return _ts + _tx + _te
                        return _wm.group(0)

                    _wtr = re.compile(r"(<w:t[^>]*>)(.*?)(</w:t>)", re.DOTALL)
                    _nh = _wtr.sub(_rdwt, _hc)
                    if _nh != _hc:
                        _hc = _nh
                        _changed = True

                    if _changed:
                        with open(_fp, "w", encoding="utf-8") as _fh:
                            _fh.write(_hc)
                        _patched += 1

        if _patched:
            os.remove(final_path)
            with _zf.ZipFile(final_path, "w", _zf.ZIP_DEFLATED) as _zout:
                for _root, _dirs, _files in os.walk(_tmpdir):
                    for _fn in _files:
                        _fp = os.path.join(_root, _fn)
                        _an = os.path.relpath(_fp, _tmpdir).replace("\\", "/")
                        _zout.write(_fp, _an)
    finally:
        _sh.rmtree(_tmpdir, ignore_errors=True)
except Exception as e:
    print(f"[ERROR] Save failed: {e}")
    sys.exit(1)

print(f"\n[DONE] {final_path}")
