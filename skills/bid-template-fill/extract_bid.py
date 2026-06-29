#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标信息提取器 - 从招标文件（PDF/DOCX/TXT/MD）提取招标方结构化信息
用法：
  python extract_bid.py <招标文件> [-o output.json] [--print]
支持格式：.pdf .docx .txt .md
"""
import json, re, sys, os
from pathlib import Path

def read_pdf(filepath):
    """从PDF提取文本"""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except ImportError:
        pass

    if not text.strip():
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except ImportError:
            pass

    if not text.strip():
        print("[WARN] 未能从PDF提取文本，请确保安装了 pdfplumber 或 PyPDF2", file=sys.stderr)
    return text


def read_docx(filepath):
    """从DOCX提取文本"""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] 需要 python-docx。请运行: pip install python-docx", file=sys.stderr)
        return ""

    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(' | '.join(cells))
    return '\n'.join(parts)


def read_file(filepath):
    """自动识别格式并读取文本"""
    ext = Path(filepath).suffix.lower()
    if ext == '.pdf':
        return read_pdf(filepath)
    elif ext == '.docx':
        return read_docx(filepath)
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()


def extract_bid_info(text):
    """从文本中提取招标信息"""
    # 去除Markdown加粗标记
    content = re.sub(r'\*\*(.+?)\*\*', r'\1', text)

    def re_first(pattern, default=""):
        m = re.search(pattern, content)
        return m.group(1).strip() if m else default

    def re_all(pattern):
        return [m.group(1).strip() for m in re.finditer(pattern, content)]

    # 招标方名称
    buyer = re_first(r'招标人[：:]\s*(.+?)(?:\n|$)')
    if not buyer:
        buyer = re_first(r'采购人[：:]\s*(.+?)(?:\n|$)')
    if not buyer:
        buyer = re_first(r'招标单位[：:]\s*(.+?)(?:\n|$)')
    if not buyer:
        buyer = re_first(r'建设单位[：:]\s*(.+?)(?:\n|$)')

    # 招标编号
    bid_no = re_first(r'(?:招标|项目)编号[：:]\s*([^\n]+)')
    if not bid_no:
        bid_no = re_first(r'(?:招标|项目)编\s*号[：:]\s*([^\n]+)')

    # 招标代理
    agency = re_first(r'招标代理[：:]\s*(.+?)(?:\n|$)')
    if not agency:
        agency = re_first(r'招标代理机构[：:]\s*(.+?)(?:\n|$)')

    # 项目名称
    project_name = re_first(r'项目名称[：:]\s*(.+?)(?:\n|$)')
    if not project_name:
        project_name = re_first(r'工程名称[：:]\s*(.+?)(?:\n|$)')

    # 标段
    lots = []
    lot_patterns = [
        r'标段\s*[：:]\s*(.+?)(?:\n|$)',
        r'(?:标段|包)\s*\d+[：:]\s*(.+?)(?:\n|$)',
        r'###?\s*\d+\.\d+\s*(.+?)(?:\n|$)',
        r'（(?:标段|包)\s*\d+）\s*(.+?)(?:\n|$)',
    ]
    for pat in lot_patterns:
        found = re_all(pat)
        if found:
            lots.extend(found)
            break

    # 截止日期
    deadline = re_first(
        r'(?:投标截止|递交截止|开标)[时间日期]?[：:]\s*(\d{4})[年/\-.](\d{1,2})[月/\-.](\d{1,2})')
    if not deadline:
        deadline = re_first(
            r'(\d{4})[年/\-.](\d{1,2})[月/\-.](\d{1,2})[日号].*?(?:截止|开标)')

    # 投标有效期
    validity = re_first(r'(?:投标有效|有效期)[：:\s]*(\d+)\s*[天日历个]')

    # 预算/限价
    budget = None
    budget_match = re.search(
        r'(?:预算|最高限价|采购预算|最高投标限价)(?:金额)?[：:\s]*(?:人民币\s*)?(\d[\d,.]*\s*(?:万元?|元))',
        content)
    if budget_match:
        raw = budget_match.group(1).replace(',', '').replace('，', '')
        _bm = re.match(r'(\d[\d.]*)\s*(.*)', raw)
        if _bm:
            budget = {
                'amount': _bm.group(1),
                'unit': _bm.group(2).strip()
            }
        else:
            budget = {'amount': raw, 'unit': ''}

    # 投标保证金
    deposit = re_first(
        r'(?:投标保证金|保证金)(?:金额)?[：:\s]*(?:人民币\s*)?(\d[\d,.]*\s*(?:万元?|元))')

    # 付款方式
    payment = re_first(r'付款方式[：:]\s*(.+?)(?:\n|$)')

    # 质保期
    warranty = re_first(r'质保期[：:]\s*(.+?)(?:\n|$)')

    # 交货期
    delivery = re_first(r'交货期[：:]\s*(.+?)(?:\n|$)')
    if not delivery:
        delivery = re_first(r'供货期[：:]\s*(.+?)(?:\n|$)')

    # 是否接受联合体
    consortium = re_first(r'联合体[：:]\s*(.+?)(?:\n|$)')

    # 分包
    subcontract = re_first(r'分包[：:]\s*(.+?)(?:\n|$)')

    # 资质要求
    qualifications = re_all(r'(?:资质要求|资格条件)[：:\s]*\d+[\.\、](.+?)(?:\n|$)')

    result = {
        'buyer': buyer,
        'bid_no': bid_no,
        'agency': agency,
        'project_name': project_name,
        'lots': lots[:5],
        'validity_days': validity,
        'deposit': deposit,
        'payment': payment,
        'warranty': warranty,
        'delivery': delivery,
        'consortium': consortium,
        'subcontract': subcontract,
        'qualifications': qualifications,
    }

    if deadline:
        result['deadline'] = {
            'year': int(deadline[0]),
            'month': int(deadline[1]),
            'day': int(deadline[2])
        }

    if budget:
        result['budget'] = budget

    return result


def main():
    if len(sys.argv) < 2:
        print("用法: python extract_bid.py <招标文件> [选项]")
        print()
        print("  支持格式: .pdf .docx .txt .md")
        print()
        print("选项:")
        print("  -o <文件>    输出为JSON文件")
        print("  --print      打印到标准输出（默认）")
        sys.exit(1)

    filepath = sys.argv[1]
    output_path = None
    print_mode = True

    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == '-o' and i + 1 < len(sys.argv):
            output_path = sys.argv[i + 1]
            print_mode = False
            i += 2
        elif sys.argv[i] == '--print':
            print_mode = True
            i += 1
        else:
            i += 1

    if not os.path.exists(filepath):
        print(f"[ERROR] 文件不存在: {filepath}", file=sys.stderr)
        sys.exit(1)

    print(f"  读取: {filepath}")
    text = read_file(filepath)

    if not text.strip():
        print("[ERROR] 未能提取到文本内容", file=sys.stderr)
        sys.exit(1)

    print(f"  文本长度: {len(text)} 字符")
    info = extract_bid_info(text)

    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(info, f, ensure_ascii=False, indent=2)
        print(f"[OK] 已保存 → {output_path}")
    else:
        print(json.dumps(info, ensure_ascii=False, indent=2))

    # 统计
    filled = sum(1 for v in info.values() if v)
    total = len(info)
    print(f"\n  提取完成: {filled}/{total} 字段已填充")


if __name__ == '__main__':
    main()
