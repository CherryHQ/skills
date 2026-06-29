#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Bid auto-generation pipeline v2.3
Parse bid doc -> extract fields -> merge company data -> fill templates -> output docx

Usage:
  python bid_pipeline.py <bid_doc> [options]

  Supported formats: .md .txt .pdf .docx

Options:
  --templates-dir <dir>      Template directory (default: bid doc directory)
  --templates <files...>     Explicit template list (space-separated)
  --company-json <file>      Company desensitized data JSON
  --company-docx <file>      Company info DOCX (auto-extract and map)
  --no-verify                Skip post-fill verification
  --output-dir <dir>         Output directory (default: bid dir /????/)

Template discovery priority:
  1. --templates explicitly specified files
  2. Template dir *desensitized*.docx files (glob match)
  3. Template dir default trio

v2.3 improvements:
  - PDF/DOCX bid doc auto-parse
  - --company-docx auto-extract from company info DOCX
  - Budget parse detects original unit, no blind *10000
  - Templates support --templates explicit + glob auto-discovery
  - Auto post-fill verification (--no-verify skip)
  - fill_large: dates-before-placeholders order

"""
import json, re, os, sys, subprocess, shutil, tempfile, zipfile, glob as globmod
from datetime import date
from pathlib import Path

try:
    from extract_company import extract_company_info as _extract_company_info
except ImportError:
    _extract_company_info = None

try:
    import pdfplumber as _pdfplumber
except ImportError:
    _pdfplumber = None
try:
    from PyPDF2 import PdfReader as _PdfReader
except ImportError:
    _PdfReader = None
try:
    from docx import Document as _Document
except ImportError:
    _Document = None

try:
    from verify_bid import verify as _verify_bid
except ImportError:
    _verify_bid = None

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

FILL_PY = Path(__file__).resolve().parent / "fill_py.py"

DEFAULT_TEMPLATE_NAMES = ["报价书_脱敏.docx", "商务标书_脱敏.docx", "技术标书_脱敏.docx"]


def read_bid_file(filepath):
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        text = ""
        if _pdfplumber:
            with _pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + chr(10)
        if not text.strip() and _PdfReader:
            reader = _PdfReader(filepath)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + chr(10)
        if not text.strip():
            raise RuntimeError("Cannot extract text from PDF. Install: pip install pdfplumber")
        return text
    elif ext == ".docx":
        if not _Document:
            raise RuntimeError("Need python-docx. Run: pip install python-docx")
        doc = _Document(filepath)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return chr(10).join(parts)
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()


def parse_bid_markdown(filepath):
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Bid file not found: {filepath}")
    raw = read_bid_file(filepath)
    if not raw.strip():
        raise ValueError(f"Bid file empty: {filepath}")
    content = re.sub("\\*\\*(.+?)\\*\\*", r"\1", raw)

    def re_first(pattern, default=""):
        m = re.search(pattern, content)
        return m.group(1).strip() if m else default

    buyer = re_first(r"招标人[：:]\s*(.+?)(?:\n|$)")
    if not buyer:
        buyer = re_first(r"采购人[：:]\s*(.+?)(?:\n|$)")
    if not buyer:
        buyer = re_first(r"招标单位[：:]\s*(.+?)(?:\n|$)")

    bid_no = re_first(r"(?:招标|项目)编号[：:]\s*([^\n]+)")
    agency = re_first(r"招标代理[：:]\s*(.+?)(?:\n|$)")
    project_name = re_first(r"项目名称[：:]\s*(.+?)(?:\n|$)")

    # Try multiple lot patterns
    lots_raw = re.findall(r"###?\s*\d+\.\d+\s*(.+?)(?:\n|$)", content)
    if not lots_raw:
        lots_raw = re.findall(r"标段\d+\s+(.+?)(?:\n|$)", content)
    if not lots_raw:
        lots_raw = re.findall(r"标段\d+\s*[|｜\s]\s*(.+?)(?:[|｜\n]|$)", content)
    lot_map = {}
    count = 0
    for name in lots_raw[:10]:
        name = re.sub(r"\s+(?:QD|CD|DHG|ISO|A\d).*$", "", name).strip()
        if 4 <= len(name) <= 40 and count < 3:
            count += 1
            lot_map[f"（标段{count}名称）"] = name

    deadline_match = re.search(
        r"(?:投标截止|递交截止|截止|开标)[时间日期]?[：:]\s*(\d{4})[年\-](\d{1,2})[月\-](\d{1,2})",
        content
    )
    deadline = None
    if deadline_match:
        deadline = {"year": int(deadline_match.group(1)), "month": int(deadline_match.group(2)), "day": int(deadline_match.group(3))}

    validity = re_first(r"(?:投标有效|有效期)[：:]?\s*(\d+)\s*[天日历]")

    budget = None
    budget_match = re.search(
        r"(?:预算|最高限价|采购预算|最高投标限价)[金额]?[：:\s]*(?:人民币\s*)?(\d[\d,.]*)\s*(万元?|元)",
        content
    )
    if budget_match:
        budget = {"amount": budget_match.group(1).replace(",", ""), "unit": budget_match.group(2)}

    deposit = re_first(r"(?:投标保证金|保证金)(?:金额)?[：:\s]*(?:人民币\s*)?(\d[\d,.]*\s*(?:万元?|元))")

    payment = re_first(r"付款方式[：:]\s*(.+?)(?:\n|$)")
    warranty = re_first(r"质保期[：:]\s*(.+?)(?:\n|$)")
    delivery = re_first(r"交货期[：:]\s*(.+?)(?:\n|$)")
    if not delivery:
        delivery = re_first(r"供货期[：:]\s*(.+?)(?:\n|$)")

    return {
        "buyer": buyer,
        "bid_no": bid_no,
        "agency": agency,
        "project_name": project_name,
        "lots": lot_map,
        "deadline": deadline,
        "validity_days": validity,
        "budget": budget,
        "deposit": deposit,
        "payment": payment,
        "warranty": warranty,
        "delivery": delivery,
    }



# ── Amount handling helpers ──

def _parse_deposit_amount(deposit_str):
    """Parse deposit string like '15万元' or '150000元' to integer"""
    import re
    m = re.match(r'(\d[\d,.]*)\s*(万元?|元)', str(deposit_str).strip())
    if m:
        val = float(m.group(1).replace(',', '').replace('，', ''))
        unit = m.group(2)
        if '万' in unit:
            val *= 10000
        return int(val)
    # Try plain number
    try:
        return int(float(str(deposit_str).replace(',', '').replace('，', '')))
    except:
        return None

def _num_to_chinese_upper(n):
    """Convert integer to Chinese uppercase (大写) like 壹拾伍万元整"""
    if n is None:
        return ""
    n = int(n)
    if n == 0:
        return "零元整"
    
    digits = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿"]
    
    def _convert_4(num):
        if num == 0:
            return ""
        s = ""
        for i in range(4):
            d = num % 10
            num //= 10
            if d > 0:
                s = digits[d] + units[i] + s
            elif s and not s.startswith("零") and num > 0:
                s = "零" + s
            if num == 0:
                break
        return s
    
    result = ""
    unit_idx = 0
    while n > 0:
        part = n % 10000
        n //= 10000
        part_str = _convert_4(part)
        if part_str:
            result = part_str + big_units[unit_idx] + result
        elif result and not result.startswith("零"):
            result = "零" + result
        unit_idx += 1
    
    return result + "元整"
def _add_amount_keys(data, deposit_num):
    """更新 data 中金额/大写相关 key 的 value 为实际金额。
    只更新已有 key 的 value，不创建新 key（确保 fill 时 key 能匹配模板文本）。"""
    import re as _re2
    amount_str = str(deposit_num)
    chinese_upper = _num_to_chinese_upper(deposit_num)

    for key in list(data.keys()):
        k = str(key)
        v = str(data[key]) if isinstance(data[key], str) else ""
        changed = v
        if _re2.search(r'￥\s*\d+[\d,]*\s*元', k):
            # key 本身含金额模式→用实际金额替换 value
            data[key] = _re2.sub(r'￥\s*\d+[\d,]*\s*元', f'￥ {amount_str} 元', k)
        if _re2.search(r'￥\s*\d+[\d,]*\s*元', v):
            changed = _re2.sub(r'￥\s*\d+[\d,]*\s*元', f'￥ {amount_str} 元', v)
        if _re2.search(r'大写\s+[零壹贰叁肆伍陆柒捌玖拾佰仟万亿元整]+', k):
            data[key] = _re2.sub(r'大写\s+[零壹贰叁肆伍陆柒捌玖拾佰仟万亿元整]+', f'大写    {chinese_upper}', k)
        if _re2.search(r'大写\s+[零壹贰叁肆伍陆柒捌玖拾佰仟万亿元整]+', v):
            changed = _re2.sub(r'大写\s+[零壹贰叁肆伍陆柒捌玖拾佰仟万亿元整]+', f'大写    {chinese_upper}', v)
        if changed != v:
            data[key] = changed

# ── End amount helpers ──

def build_fill_json(info, company_data=None):
    data = {}
    if company_data:
        data.update(company_data)

    if info.get("project_name"):
        data["[项目名称]"] = info["project_name"]
        data["_[项目名称]_"] = info["project_name"]
    if info.get("bid_no"):
        data["_[招标编号]_"] = info["bid_no"]
        data["[招标编号]"] = info["bid_no"]
    if info.get("buyer"):
        data["_[招标方名称]_"] = info["buyer"]
        data["[招标方名称]"] = info["buyer"]
    if info.get("agency"):
        data["_[招标代理名称]_"] = info["agency"]
        data["[招标代理名称]"] = info["agency"]

    lots = info.get("lots", {})
    for placeholder, name in lots.items():
        data[placeholder] = name

    if info.get("validity_days"):
        days = info["validity_days"]
        data["【 30 】"] = f"【 {days} 】"

    # Parse deposit for amount replacement
    deposit_str = info.get("deposit", "")
    deposit_num = None
    if deposit_str:
        deposit_num = _parse_deposit_amount(deposit_str)

    if info.get("deposit"):
        data["【 】"] = f"【 人民币 {info['deposit']} 】"

    # Add amount replacement keys for desensitized template amounts
    if deposit_num:
        _add_amount_keys(data, deposit_num)

    return data


def load_company_data(templates_dir, company_json_path):
    if company_json_path and os.path.exists(company_json_path):
        with open(company_json_path, "r", encoding="utf-8-sig") as f:
            data = json.load(f)
        # Handle nested format from extract_company.py: {"fill_data": {...}, ...}
        if isinstance(data, dict) and "fill_data" in data:
            return data["fill_data"]
        return data
    if templates_dir:
        default = os.path.join(templates_dir, "company.json")
        if os.path.exists(default):
            with open(default, "r", encoding="utf-8-sig") as f:
                data = json.load(f)
            if isinstance(data, dict) and "fill_data" in data:
                return data["fill_data"]
            return data
    return {}


def discover_templates(templates_dir, explicit=None):
    templates = []
    if explicit:
        for tp in explicit:
            tp_path = Path(tp)
            if not tp_path.is_absolute():
                tp_path = Path(templates_dir) / tp_path if templates_dir else Path.cwd() / tp_path
            if tp_path.exists() and tp_path.suffix.lower() == ".docx":
                templates.append(tp_path)
        if templates:
            return templates

    if templates_dir:
        td = Path(templates_dir)
        if td.is_dir():
            for pattern in ["*脱敏*.docx"]:
                for fpath in td.glob(pattern):
                    if fpath not in templates:
                        templates.append(fpath)
            if templates:
                return templates
            for name in DEFAULT_TEMPLATE_NAMES:
                fp = td / name
                if fp.exists() and fp not in templates:
                    templates.append(fp)

    return templates


def fill_small(template_path, output_path, json_path):
    try:
        result = subprocess.run(
            [sys.executable, str(FILL_PY), template_path, output_path, json_path],
            capture_output=True, text=True, timeout=300
        )
        if result.returncode == 0:
            print(result.stdout.strip()[-200:] if len(result.stdout.strip()) > 200 else result.stdout.strip())
            return True
        else:
            print(result.stderr.strip()[:500] if result.stderr else "Unknown error")
            return False
    except Exception as e:
        print(f"  [ERROR] fill_small failed: {e}")
        return False


def fill_large(template_path, output_path, data):
    """Pure Python zipfile fill for large files (5MB+). Zero external deps.
    Phase 1: dates -> Phase 2: placeholders. Handles headers/footers/body."""
    today = date.today()
    today_str = f"{today.year} 年 {today.month:02d} 月 {today.day:02d} 日"
    tmp = tempfile.mkdtemp(prefix=".tmp_bid_")
    try:
        with zipfile.ZipFile(template_path, "r") as zf:
            zf.extractall(tmp)

        sorted_keys = sorted(data.items(), key=lambda x: len(x[0]), reverse=True)

        import re as _re

        for root, dirs, files in os.walk(tmp):
            for fname in files:
                if not fname.endswith(".xml") and not fname.endswith(".rels"):
                    continue
                fpath = os.path.join(root, fname)
                with open(fpath, "r", encoding="utf-8") as f:
                    fcontent = f.read()

                changed = False
                rel_path = os.path.relpath(fpath, tmp).replace("\\", "/")

                # Phase 1: Date replacement FIRST (before placeholder values interfere)
                if rel_path.startswith("word/"):
                    d1 = _re.compile(r"(20\d{2})[_\s]*年[_\s]*(\d{1,2})[_\s]*月[_\s]*(\d{1,2})[_\s]*日")
                    d2 = _re.compile(r"_*(20\d{2})[_\s]*年[_\s]*月[_\s]*日")
                    d3 = _re.compile(r"_*年[_\s]*月[_\s]*日")
                    def dr(m):
                        y = int(m.group(1))
                        mo = int(m.group(2))
                        d = int(m.group(3))
                        if y == today.year and mo == today.month and d == today.day:
                            return m.group(0)
                        return today_str
                    nt, c1 = d1.subn(dr, fcontent)
                    nt, c2 = d2.subn(today_str, nt)
                    nt, c3 = d3.subn(today_str, nt)
                    if c1 + c2 + c3 > 0:
                        fcontent = nt
                        changed = True

                # Phase 2: XML-aware placeholder replacement for word/ files, plain text for others
                if rel_path.startswith("word/") and fname.endswith(".xml"):
                    wt_re = _re.compile(r"(<w:t[^>]*>)(.*?)(</w:t>)", _re.DOTALL)
                    # Merge <w:t> runs per paragraph for XML-aware replacement
                    def _replace_para(p_match):
                        para_text = p_match.group(0)
                        # Collect all <w:t> text
                        texts = []
                        def _collect(m):
                            texts.append(m.group(2))
                            return m.group(0)
                        wt_re.sub(_collect, para_text)
                        merged = "".join(texts)
                        if not merged.strip():
                            return para_text
                        orig = merged
                        for key, value in sorted_keys:
                            if key and key in merged:
                                escaped = str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\"", "&quot;")
                                merged = merged.replace(key, escaped)
                        if merged == orig:
                            return para_text
                        # Rebuild: put merged text in first <w:t>, clear others
                        first = [True]
                        def _rebuild(m):
                            if first[0]:
                                first[0] = False
                                return m.group(1) + merged + m.group(3)
                            return m.group(1) + m.group(3)
                        return wt_re.sub(_rebuild, para_text)
                    p_re = _re.compile(r"<w:p\b[\s\S]*?</w:p>", _re.DOTALL)
                    new_content = p_re.sub(_replace_para, fcontent)
                    if new_content != fcontent:
                        fcontent = new_content
                        changed = True
                else:
                    for key, value in sorted_keys:
                        if not key:
                            continue
                        if key in fcontent:
                            escaped = str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\"", "&quot;")
                            fcontent = fcontent.replace(key, escaped)
                            changed = True

                if changed:
                    with open(fpath, "w", encoding="utf-8") as f:
                        f.write(fcontent)

        zip_path = output_path.replace(".docx", ".zip")
        for p in [zip_path, output_path]:
            if os.path.exists(p):
                os.remove(p)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            ct_path = os.path.join(tmp, "[Content_Types].xml")
            if os.path.exists(ct_path):
                zf.write(ct_path, "[Content_Types].xml")
            for root, dirs, files in os.walk(tmp):
                for fname in files:
                    fpath = os.path.join(root, fname)
                    arcname = os.path.relpath(fpath, tmp).replace("\\", "/")
                    if arcname == "[Content_Types].xml":
                        continue
                    zf.write(fpath, arcname)

        if os.path.exists(zip_path):
            os.rename(zip_path, output_path)
        return True
    except Exception as e:
        print(f"  [ERROR] Large file fill failed: {e}")
        return False
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def main():
    if len(sys.argv) < 2:
        print("Usage: python bid_pipeline.py <bid_doc> [options]")
        print("  --templates-dir <dir>   Template directory")
        print("  --templates <files...>  Explicit templates")
        print("  --company-json <file>   Company JSON data")
        print("  --company-docx <file>   Company info DOCX")
        print("  --no-verify             Skip post-fill check")
        print("  --output-dir <dir>      Output directory")
        sys.exit(1)

    bid_file = sys.argv[1]
    if not os.path.exists(bid_file):
        print(f"[ERROR] Bid file not found: {bid_file}")
        sys.exit(1)

    templates_dir = str(Path(bid_file).parent)
    explicit_templates = None
    company_json_path = None
    company_docx_path = None
    output_dir = None
    _skip_verify = False

    i = 2
    while i < len(sys.argv):
        arg = sys.argv[i]
        if arg == "--templates-dir" and i + 1 < len(sys.argv):
            templates_dir = sys.argv[i + 1]
            i += 2
        elif arg == "--templates":
            explicit_templates = []
            i += 1
            while i < len(sys.argv) and not sys.argv[i].startswith("--"):
                explicit_templates.append(sys.argv[i])
                i += 1
        elif arg == "--company-json" and i + 1 < len(sys.argv):
            company_json_path = sys.argv[i + 1]
            i += 2
        elif arg == "--company-docx" and i + 1 < len(sys.argv):
            company_docx_path = sys.argv[i + 1]
            i += 2
        elif arg == "--output-dir" and i + 1 < len(sys.argv):
            output_dir = sys.argv[i + 1]
            i += 2
        elif arg == "--no-verify":
            _skip_verify = True
            i += 1
        else:
            i += 1

    print("=" * 50)
    print("  Bid Auto-Generation Pipeline v2.3")
    print("=" * 50)

    company_data = load_company_data(templates_dir, company_json_path)
    if not company_data and company_docx_path:
        if _extract_company_info:
            print(f"\n  Extracting company info from: {company_docx_path}")
            extracted = _extract_company_info(company_docx_path)
            company_data = extracted.get("fill_data", {})
            if company_data:
                print(f"  Auto-mapped {len(company_data)} template keys")
        else:
            print("  [WARN] extract_company.py not found, skipping DOCX extraction")

    if company_data:
        print(f"\n  Company data: {len(company_data)} keys")
    else:
        print("\n  [WARN] No company data found, filling bid-side info only")

    ext = Path(bid_file).suffix.lower()
    print(f"\n[1/4] Parsing bid doc: {bid_file} ({ext})")
    try:
        info = parse_bid_markdown(bid_file)
    except (FileNotFoundError, ValueError) as e:
        print(f"  [ERROR] {e}")
        sys.exit(1)

    print(f"  Buyer: {info['buyer'] or '(not found)'}")
    print(f"  Project: {info['project_name'] or '(not found)'}")
    print(f"  Bid No: {info['bid_no'] or '(not found)'}")
    print(f"  Agency: {info['agency'] or '(not found)'}")

    print(f"\n[2/4] Building fill data...")
    data = build_fill_json(info, company_data)
    json_path = os.path.join(os.path.dirname(bid_file), ".bid_pipeline_data.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"  Data written: {json_path} ({len(data)} keys)")

    templates = discover_templates(templates_dir, explicit_templates)
    if not templates:
        print("\n  [ERROR] No templates found")
        if os.path.exists(json_path):
            os.remove(json_path)
        sys.exit(1)

    print(f"\n[3/4] Found {len(templates)} templates:")
    for t in templates:
        print(f"  - {t.name}")

    out_dir = Path(output_dir) if output_dir else Path(bid_file).parent / "成品标书"
    out_dir.mkdir(parents=True, exist_ok=True)
    all_ok = True
    for tmpl_path in templates:
        stem = tmpl_path.stem.replace("_脱敏", "")
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", (info["project_name"] or "output")[:15])
        out_name = f"{stem}_{safe_name}.docx"
        out_path = out_dir / out_name

        print(f"\n[4/4] Filling: {tmpl_path.name} => {out_name}")
        size_mb = os.path.getsize(str(tmpl_path)) / 1024 / 1024
        if size_mb > 5:
            ok = fill_large(str(tmpl_path), str(out_path), data)
        else:
            ok = fill_small(str(tmpl_path), str(out_path), json_path)
        if ok:
            out_mb = os.path.getsize(str(out_path)) / 1024 / 1024
            print(f"  [OK] {out_name} ({out_mb:.1f} MB)")
        else:
            print(f"  [FAIL] {tmpl_path.name}")
            all_ok = False

    # Post-fill verification
    if not _skip_verify and _verify_bid:
        print(f"\n[5/5] Verifying...")
        for tmpl_path in templates:
            stem = tmpl_path.stem.replace("_脱敏", "")
            safe_name = re.sub(r'[\\/:*?"<>|]', "_", (info["project_name"] or "output")[:15])
            out_name = f"{stem}_{safe_name}.docx"
            out_path = out_dir / out_name
            if out_path.exists():
                report = _verify_bid(str(out_path), company_docx_path, strict=False)
                p0 = report.get("p0_fatal", 0)
                p1 = report.get("p1_serious", 0)
                if p0 > 0:
                    print(f"  [FAIL] {out_name} - {p0} fatal issues")
                    all_ok = False
                elif p1 > 0:
                    print(f"  [WARN] {out_name} - {p1} serious issues")
                else:
                    print(f"  [OK] {out_name} - passed")
    elif not _verify_bid:
        print(f"\n  [INFO] verify_bid.py not found, skipping verification")

    if os.path.exists(json_path):
        os.remove(json_path)

    status = "All done" if all_ok else "Partial failure"
    print(f"\n[DONE] {status} - output: {out_dir}")

if __name__ == "__main__":
    main()
