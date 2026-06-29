#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公司信息提取器 - 从公司内部信息DOCX提取结构化数据，输出为company.json格式
用法：
  python extract_company.py <公司信息.docx> [-o company.json]
  python extract_company.py <公司信息.docx> --print
  python extract_company.py <公司信息.docx> --map desensitize_map.json

映射文件 (desensitize_map.json) 定义模板脱敏文本与公司字段的对应关系。
第一次使用前，请编辑 desensitize_map.json 中的 template_desensitize_map，
将左侧的虚拟文本改为你模板中实际出现的脱敏文本。
"""
import json, re, sys, os
from pathlib import Path

def _get_first_value(raw_fields, variant_list):
    """从 raw_fields 中按变体列表顺序查找第一个非空值"""
    for name in variant_list:
        val = raw_fields.get(name, '')
        if val:
            return val
    return ''

def _load_map(map_path):
    """加载脱敏映射文件"""
    if map_path and os.path.exists(map_path):
        with open(map_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    # 尝试默认位置
    default = Path(__file__).resolve().parent / 'desensitize_map.json'
    if default.exists():
        with open(str(default), 'r', encoding='utf-8') as f:
            return json.load(f)
    return None


def extract_company_info(docx_path, map_path=None):
    """从公司信息DOCX提取结构化数据，根据映射文件匹配模板key"""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] 需要 python-docx。请运行: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(docx_path):
        print(f"[ERROR] 文件不存在: {docx_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(docx_path)

    all_texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_texts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            all_texts.append('\t'.join(cells))

    full_text = '\n'.join(all_texts)

    field_pattern = re.compile(
        r'^(.*?)[：:]\s*(.+?)$', re.MULTILINE
    )

    raw_fields = {}
    for m in field_pattern.finditer(full_text):
        key = m.group(1).strip()
        value = m.group(2).strip()
        if len(key) > 20 or key in ('一', '二', '三', '四', '五', '六', '七', '八', '附'):
            continue
        if not value or value in ('—', '-'):
            continue
        raw_fields[key] = value

    # 加载映射配置
    cfg = _load_map(map_path) or {}

    # 按变体列表提取公司字段
    def _fv(section, default_val=''):
        variants = (cfg.get(section, {}) or {}).get('values', [])
        return _get_first_value(raw_fields, variants) or default_val

    COMPANY_FULL = _fv('company_name_variants')
    COMPANY_SHORT = _fv('company_short_variants')
    LEGAL_PERSON = _fv('legal_person_variants')
    AUTH_REP = _fv('auth_rep_variants')
    ADDRESS = _fv('address_variants')
    ZIPCODE = _fv('zipcode_variants')
    BANK = _fv('bank_variants')
    ACCOUNT = _fv('account_variants')
    PHONE = _fv('phone_variants')
    MOBILE = _fv('mobile_variants')
    FAX = _fv('fax_variants')
    CONTACT = _fv('contact_variants')
    credit_code = _fv('credit_code_variants')

    result = {}

    # 根据 template_desensitize_map 映射：模板脱敏文本 → 公司实际值
    des_map = cfg.get('template_desensitize_map', {})
    if des_map:
        field_values = {
            '公司全称': COMPANY_FULL,
            '公司简称': COMPANY_SHORT,
            '法定代表人': LEGAL_PERSON,
            '授权代表': AUTH_REP,
            '注册地址': ADDRESS,
            '邮政编码': ZIPCODE,
            '开户银行全称': BANK,
            '银行账号': ACCOUNT,
            '联系电话': PHONE,
            '手机号码': MOBILE,
            '传真号码': FAX,
        }
        for template_text, field_name in des_map.items():
            actual_value = field_values.get(field_name, raw_fields.get(field_name, ''))
            if actual_value:
                result[template_text] = actual_value

    # 附加信息
    extra = {}
    if CONTACT:
        extra['联系人'] = CONTACT
    if credit_code:
        extra['统一社会信用代码'] = credit_code

    return {'fill_data': result, 'extra': extra, 'raw': raw_fields}


def main():
    if len(sys.argv) < 2:
        print("用法: python extract_company.py <公司信息.docx> [选项]")
        print()
        print("选项:")
        print("  -o <文件>      输出为JSON文件")
        print("  --print        打印到标准输出（默认）")
        print("  --map <文件>   指定脱敏映射文件（默认: desensitize_map.json）")
        print()
        print("示例:")
        print("  python extract_company.py 公司信息.docx -o company.json")
        print("  python extract_company.py 公司信息.docx --print")
        print("  python extract_company.py 公司信息.docx --map my_map.json")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_path = None
    print_mode = True
    map_path = None

    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == '-o' and i + 1 < len(sys.argv):
            output_path = sys.argv[i + 1]
            print_mode = False
            i += 2
        elif sys.argv[i] == '--print':
            print_mode = True
            i += 1
        elif sys.argv[i] == '--map' and i + 1 < len(sys.argv):
            map_path = sys.argv[i + 1]
            i += 2
        else:
            i += 1

    data = extract_company_info(docx_path, map_path=map_path)

    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data['fill_data'], f, ensure_ascii=False, indent=2)
        print(f"[OK] 已提取 {len(data['fill_data'])} 个映射键 → {output_path}")
        if data['extra']:
            print(f"  附加信息: {json.dumps(data['extra'], ensure_ascii=False)}")
    else:
        print("=" * 50)
        print("  公司信息提取结果")
        print("=" * 50)
        print()
        print("--- 模板填充映射（可直接用作company.json）---")
        print(json.dumps(data['fill_data'], ensure_ascii=False, indent=2))
        if data['extra']:
            print()
            print("--- 附加参考信息 ---")
            print(json.dumps(data['extra'], ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
